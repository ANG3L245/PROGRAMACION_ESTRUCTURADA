import re
import os
import funciones
from finanzas import crud_finanzas
from finanzas import graficas
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
from reportlab.lib import colors
import openpyxl
from docx import Document

# Variables globales para colores
COLOR_BORDE = "\033[1;34m"
COLOR_TITULO = "\033[1;36m"
COLOR_TEXTO = "\033[1;30m"
COLOR_SALIR = "\033[1;31m"
COLOR_ALERTA = "\033[1;33m"
RESET = "\033[0m"

CATEGORIAS = {
    "gasto": ["COMIDA", "TRANSPORTE", "ENTRETENIMIENTO", "SERVICIOS", "OTROS"],
    "ingreso": ["SUELDO", "BECA", "VENTA", "REGALO", "OTROS"]
}

def menuFinanzas(conexionBD, usuario_activo):
    opc = ""
    while opc != "8":
        try:
            funciones.borrarPantalla()
            
            titulo = f"::::: MENU FINANZAS ({usuario_activo}) :::::"
            
            print(f"\n\t{COLOR_BORDE}╔═══════════════════════════════════════════════════════╗{RESET}")
            print(f"\t{COLOR_BORDE}║{RESET}{COLOR_TITULO}{titulo:^55}{RESET}{COLOR_BORDE}║{RESET}")
            print(f"\t{COLOR_BORDE}╠═══════════════════════════════════════════════════════╣{RESET}")
            print(f"\t{COLOR_BORDE}║{RESET}                                                       {COLOR_BORDE}║{RESET}")
            print(f"\t{COLOR_BORDE}║{RESET}  {COLOR_TEXTO}1.- Agregar movimiento{RESET}                                {COLOR_BORDE}║{RESET}")
            print(f"\t{COLOR_BORDE}║{RESET}  {COLOR_TEXTO}2.- Ver movimientos{RESET}                                   {COLOR_BORDE}║{RESET}")
            print(f"\t{COLOR_BORDE}║{RESET}  {COLOR_TEXTO}3.- Editar / Eliminar movimiento{RESET}                      {COLOR_BORDE}║{RESET}")
            print(f"\t{COLOR_BORDE}║{RESET}  {COLOR_TEXTO}4.- Ver balance{RESET}                                       {COLOR_BORDE}║{RESET}")
            print(f"\t{COLOR_BORDE}║{RESET}  {COLOR_TEXTO}5.- Ver grafica de gastos{RESET}                             {COLOR_BORDE}║{RESET}")
            print(f"\t{COLOR_BORDE}║{RESET}  {COLOR_TEXTO}6.- Guardar grafica de gastos como imagen{RESET}             {COLOR_BORDE}║{RESET}")
            print(f"\t{COLOR_BORDE}║{RESET}  {COLOR_TEXTO}7.- Descargar movimientos (PDF / Excel / Word){RESET}        {COLOR_BORDE}║{RESET}")
            print(f"\t{COLOR_BORDE}║{RESET}  {COLOR_SALIR}8.- Volver al menu principal{RESET}                          {COLOR_BORDE}║{RESET}")
            print(f"\t{COLOR_BORDE}║{RESET}                                                       {COLOR_BORDE}║{RESET}")
            print(f"\t{COLOR_BORDE}╚═══════════════════════════════════════════════════════╝{RESET}\n")

            opc = input(f"\t{COLOR_TEXTO}Escribe una opcion:{RESET} ").strip()

            match opc:
                case "1":
                    funciones.borrarPantalla()
                    agregarMovimiento(conexionBD, usuario_activo)
                case "2":
                    funciones.borrarPantalla()
                    verMovimientos(conexionBD, usuario_activo)
                case "3":
                    funciones.borrarPantalla()
                    gestionarMovimiento(conexionBD, usuario_activo)
                case "4":
                    funciones.borrarPantalla()
                    verBalance(conexionBD, usuario_activo)
                case "5":
                    funciones.borrarPantalla()
                    mostrarGrafica(conexionBD, usuario_activo)
                case "6":
                    funciones.borrarPantalla()
                    exportarGrafica(conexionBD, usuario_activo)
                case "7":
                    funciones.borrarPantalla()
                    exportarMovimientos(conexionBD, usuario_activo)
                case "8":
                    pass
                case _:
                    funciones.opcionInvalida()
                    
        except KeyboardInterrupt:
            print(f"\n\t{COLOR_ALERTA}...Operacion cancelada por el usuario (Ctrl+C)...{RESET}")
            funciones.esperarTecla()
        except Exception as e:
            print(f"\n\t{COLOR_ALERTA}...Ocurrio un error inesperado: {e}...{RESET}")
            funciones.esperarTecla()


def elegirCategoria(tipo):
    lista = CATEGORIAS[tipo]
    print(f"\n{COLOR_TITULO}Categorias Disponibles:{RESET}")
    for i, cat in enumerate(lista, start=1):
        print(f"\t{i}. {cat}")
    opc = input(f"{COLOR_TEXTO}Elige el numero de categoria:{RESET} ").strip()
    try:
        return lista[int(opc) - 1]
    except (ValueError, IndexError):
        return "OTROS"


def agregarMovimiento(conexionBD, usuario_activo):
    try:
        print(f"\n\t\t{COLOR_TITULO}::::: AGREGAR MOVIMIENTO :::::{RESET}\n")
        tipo = input(f"{COLOR_TEXTO}Tipo (ingreso/gasto):{RESET} ").lower().strip()
        while tipo not in ("ingreso", "gasto"):
            print(f"\t{COLOR_ALERTA}...Tipo invalido, escribe 'ingreso' o 'gasto'...{RESET}")
            tipo = input(f"{COLOR_TEXTO}Tipo (ingreso/gasto):{RESET} ").lower().strip()

        categoria = elegirCategoria(tipo)
        
        monto = input(f"{COLOR_TEXTO}Monto: ${RESET}").strip()
        while not re.match(r"^\d+(\.\d{1,2})?$", monto):
            print(f"\n\t{COLOR_ALERTA}...Monto invalido. Ingrese un valor numerico positivo (ej. 150.50)...{RESET}")
            monto = input(f"{COLOR_TEXTO}Monto: ${RESET}").strip()

        descripcion = input(f"{COLOR_TEXTO}Descripcion (opcional):{RESET} ").strip()
        
        fecha = input(f"{COLOR_TEXTO}Fecha (YYYY-MM-DD):{RESET} ").strip()
        while not re.match(r"^\d{4}-\d{2}-\d{2}$", fecha):
            print(f"\n\t{COLOR_ALERTA}...Formato de fecha invalido. Use el formato YYYY-MM-DD...{RESET}")
            fecha = input(f"{COLOR_TEXTO}Fecha (YYYY-MM-DD):{RESET} ").strip()

        respuesta = crud_finanzas.insertar(usuario_activo, tipo, categoria, monto, descripcion, fecha, conexionBD)
        if respuesta:
            funciones.accionExitosa()
        else:
            funciones.accionNoExitosa()
            
    except Exception as e:
        print(f"\n\t{COLOR_ALERTA}...Error al agregar movimiento: {e}...{RESET}")
        funciones.esperarTecla()


def verMovimientos(conexionBD, usuario_activo):
    try:
        print(f"\n\t\t{COLOR_TITULO}::::: MOVIMIENTOS :::::{RESET}\n")
        movimientos = crud_finanzas.consultarPorUsuario(usuario_activo, conexionBD)

        if len(movimientos) > 0:
            print(f"\t{COLOR_TITULO}{'ID':<5}\t{'Tipo':<10}\t{'Categoria':<18}\t{'Monto':>10}\t{'Fecha':<12}{RESET}\n")
            for i in movimientos:
                print(f"\t{COLOR_TEXTO}{i[0]:<5}\t{i[2]:<10}\t{i[3]:<18}\t{float(i[4]):>10.2f}\t{str(i[6]):<12}{RESET}")
            print(f"\t{COLOR_BORDE}" + "-" * 90 + f"{RESET}")
            funciones.esperarTecla()
        else:
            print(f"\n\t{COLOR_ALERTA}...No hay movimientos registrados, verifique!...{RESET}")
            funciones.esperarTecla()
            
    except Exception as e:
        print(f"\n\t{COLOR_ALERTA}...Error al consultar movimientos: {e}...{RESET}")
        funciones.esperarTecla()


def gestionarMovimiento(conexionBD, usuario_activo):
    try:
        movimientos = crud_finanzas.consultarPorUsuario(usuario_activo, conexionBD)

        if len(movimientos) == 0:
            print(f"\n\t{COLOR_ALERTA}...No hay movimientos registrados, verifique!...{RESET}")
            funciones.esperarTecla()
            return

        print(f"\n\t\t{COLOR_TITULO}::::: TUS MOVIMIENTOS :::::{RESET}\n")
        print(f"\t{COLOR_TITULO}{'ID':<5}\t{'Tipo':<10}\t{'Categoria':<18}\t{'Monto':>10}\t{'Fecha':<12}{RESET}\n")

        for i in movimientos:
            print(f"\t{COLOR_TEXTO}{i[0]:<5}\t{i[2]:<10}\t{i[3]:<18}\t{float(i[4]):>10.2f}\t{str(i[6]):<12}{RESET}")
        print(f"\t{COLOR_BORDE}" + "-" * 90 + f"{RESET}")

        id_movimiento = input(f"\n{COLOR_TEXTO}Escribe el ID del movimiento a editar/eliminar:{RESET} ").strip()

        ids_validos = [str(m[0]) for m in movimientos]
        if id_movimiento not in ids_validos:
            print(f"\n\t{COLOR_ALERTA}...Ese ID no pertenece a ninguno de tus movimientos...{RESET}")
            funciones.esperarTecla()
            return

        print(f"\n\t{COLOR_TEXTO}1.- Editar{RESET}")
        print(f"\t{COLOR_TEXTO}2.- Eliminar{RESET}")
        accion = input(f"{COLOR_TEXTO}Elige una opcion:{RESET} ").strip()

        if accion == "1":
            editarMovimiento(conexionBD, id_movimiento)
        elif accion == "2":
            eliminarMovimiento(conexionBD, id_movimiento)
        else:
            funciones.opcionInvalida()
            
    except Exception as e:
        print(f"\n\t{COLOR_ALERTA}...Error al gestionar movimiento: {e}...{RESET}")
        funciones.esperarTecla()


def editarMovimiento(conexionBD, id_movimiento):
    try:
        print(f"\n\t\t{COLOR_TITULO}::::: EDITAR MOVIMIENTO :::::{RESET}\n")
        tipo = input(f"{COLOR_TEXTO}Nuevo tipo (ingreso/gasto):{RESET} ").lower().strip()
        while tipo not in ("ingreso", "gasto"):
            print(f"\t{COLOR_ALERTA}...Tipo invalido, escribe 'ingreso' o 'gasto'...{RESET}")
            tipo = input(f"{COLOR_TEXTO}Nuevo tipo (ingreso/gasto):{RESET} ").lower().strip()

        categoria = elegirCategoria(tipo)
        
        monto = input(f"{COLOR_TEXTO}Nuevo monto: ${RESET}").strip()
        while not re.match(r"^\d+(\.\d{1,2})?$", monto):
            print(f"\n\t{COLOR_ALERTA}...Monto invalido. Ingrese un valor numerico positivo (ej. 150.50)...{RESET}")
            monto = input(f"{COLOR_TEXTO}Nuevo monto: ${RESET}").strip()

        descripcion = input(f"{COLOR_TEXTO}Nueva descripcion (opcional):{RESET} ").strip()
        
        fecha = input(f"{COLOR_TEXTO}Nueva fecha (YYYY-MM-DD):{RESET} ").strip()
        while not re.match(r"^\d{4}-\d{2}-\d{2}$", fecha):
            print(f"\n\t{COLOR_ALERTA}...Formato de fecha invalido. Use el formato YYYY-MM-DD...{RESET}")
            fecha = input(f"{COLOR_TEXTO}Nueva fecha (YYYY-MM-DD):{RESET} ").strip()

        respuesta = crud_finanzas.actualizar(id_movimiento, tipo, categoria, monto, descripcion, fecha, conexionBD)
        if respuesta:
            funciones.accionExitosa()
        else:
            funciones.accionNoExitosa()
            
    except Exception as e:
        print(f"\n\t{COLOR_ALERTA}...Error al editar movimiento: {e}...{RESET}")
        funciones.esperarTecla()


def eliminarMovimiento(conexionBD, id_movimiento):
    try:
        opc = input(f"\n{COLOR_TEXTO}¿Seguro que deseas eliminar el movimiento {id_movimiento}? (si/no):{RESET} ").lower().strip()

        if opc == "si":
            respuesta = crud_finanzas.eliminar(id_movimiento, conexionBD)
            if respuesta:
                funciones.accionExitosa()
            else:
                funciones.accionNoExitosa()
        else:
            print(f"\n\t{COLOR_ALERTA}...Operacion cancelada...{RESET}")
            funciones.esperarTecla()
            
    except Exception as e:
        print(f"\n\t{COLOR_ALERTA}...Error al eliminar movimiento: {e}...{RESET}")
        funciones.esperarTecla()


def verBalance(conexionBD, usuario_activo):
    try:
        movimientos = crud_finanzas.consultarPorUsuario(usuario_activo, conexionBD)

        ingresos = sum(float(m[4]) for m in movimientos if m[2] == "ingreso")
        gastos = sum(float(m[4]) for m in movimientos if m[2] == "gasto")
        balance = ingresos - gastos

        print(f"\n\t\t{COLOR_TITULO}::::: BALANCE :::::{RESET}\n")
        print(f"\t{COLOR_TEXTO}Total ingresos: ${ingresos:.2f}{RESET}")
        print(f"\t{COLOR_TEXTO}Total gastos:   ${gastos:.2f}{RESET}")
        print(f"\t{COLOR_TEXTO}Balance:        ${balance:.2f}{RESET}")
        funciones.esperarTecla()
        
    except Exception as e:
        print(f"\n\t{COLOR_ALERTA}...Error al calcular el balance: {e}...{RESET}")
        funciones.esperarTecla()


def mostrarGrafica(conexionBD, usuario_activo):
    try:
        datos = crud_finanzas.obtenerGastosPorCategoria(usuario_activo, conexionBD)

        if len(datos) > 0:
            categorias = [d[0] for d in datos]
            montos = [float(d[1]) for d in datos]
            graficas.verGrafica(categorias, montos)
        else:
            print(f"\n\t{COLOR_ALERTA}...No hay gastos registrados para graficar!...{RESET}")
            funciones.esperarTecla()
            
    except Exception as e:
        print(f"\n\t{COLOR_ALERTA}...Error al mostrar la grafica: {e}...{RESET}")
        funciones.esperarTecla()


def exportarGrafica(conexionBD, usuario_activo):
    try:
        datos = crud_finanzas.obtenerGastosPorCategoria(usuario_activo, conexionBD)

        if len(datos) > 0:
            categorias = [d[0] for d in datos]
            montos = [float(d[1]) for d in datos]
            graficas.guardarGrafica(categorias, montos)
            funciones.esperarTecla()
        else:
            print(f"\n\t{COLOR_ALERTA}...No hay gastos registrados para graficar!...{RESET}")
            funciones.esperarTecla()
            
    except Exception as e:
        print(f"\n\t{COLOR_ALERTA}...Error al exportar la grafica: {e}...{RESET}")
        funciones.esperarTecla()


def exportarMovimientos(conexionBD, usuario_activo):
    try:
        movimientos = crud_finanzas.consultarPorUsuario(usuario_activo, conexionBD)

        if len(movimientos) == 0:
            print(f"\n\t{COLOR_ALERTA}...No hay movimientos registrados para exportar!...{RESET}")
            funciones.esperarTecla()
            return

        print(f"\n\t\t{COLOR_TITULO}::::: FORMATO DE DESCARGA :::::{RESET}\n")
        print(f"\t{COLOR_TEXTO}1.- Documento PDF (.pdf){RESET}")
        print(f"\t{COLOR_TEXTO}2.- Hoja de Excel (.xlsx){RESET}")
        print(f"\t{COLOR_TEXTO}3.- Documento Word (.docx){RESET}")
        
        opc_formato = input(f"\n\t{COLOR_TEXTO}Elige el numero de formato (1/2/3):{RESET} ").strip()
        if opc_formato not in ("1", "2", "3"):
            print(f"\n\t{COLOR_ALERTA}...Opcion invalida...{RESET}")
            funciones.esperarTecla()
            return

        ruta_carpeta = "/home/zwain/proyecto_final1/archivos/"
        os.makedirs(ruta_carpeta, exist_ok=True) 

        encabezados = ["ID", "Tipo", "Categoria", "Monto", "Fecha"]

        if opc_formato == "1":
            ruta = os.path.join(ruta_carpeta, f"movimientos_{usuario_activo}.pdf")
            datos = [encabezados]
            for i in movimientos:
                datos.append([str(i[0]), i[2], i[3], f"${float(i[4]):.2f}", str(i[6])])

            doc = SimpleDocTemplate(ruta, pagesize=letter)
            tabla = Table(datos)
            tabla.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
            ]))
            doc.build([tabla])

        elif opc_formato == "2":
            ruta = os.path.join(ruta_carpeta, f"movimientos_{usuario_activo}.xlsx")
            libro = openpyxl.Workbook()
            hoja = libro.active
            hoja.title = "Mis Movimientos"
            hoja.append(encabezados)
            
            for i in movimientos:
                hoja.append([i[0], i[2], i[3], float(i[4]), str(i[6])])
            libro.save(ruta)

        elif opc_formato == "3":
            ruta = os.path.join(ruta_carpeta, f"movimientos_{usuario_activo}.docx")
            doc = Document()
            doc.add_heading(f'Reporte de Movimientos - {usuario_activo}', 0)
            
            tabla = doc.add_table(rows=1, cols=len(encabezados))
            tabla.style = 'Table Grid'
            
            hdr_cells = tabla.rows[0].cells
            for idx, nombre_columna in enumerate(encabezados):
                hdr_cells[idx].text = nombre_columna
                
            for i in movimientos:
                row_cells = tabla.add_row().cells
                row_cells[0].text = str(i[0])
                row_cells[1].text = str(i[2])
                row_cells[2].text = str(i[3])
                row_cells[3].text = f"${float(i[4]):.2f}"
                row_cells[4].text = str(i[6])
                
            doc.save(ruta)

        print(f"\n\t{COLOR_TITULO}...Archivo exportado exitosamente en:\n\t{ruta}{RESET}")
        funciones.esperarTecla()
        
    except Exception as e:
        print(f"\n\t{COLOR_ALERTA}...Error al exportar el archivo: {e}...{RESET}")
        funciones.esperarTecla()